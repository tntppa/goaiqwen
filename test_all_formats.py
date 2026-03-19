"""
全格式端到端测试：PDF、图片、XLS/XLSX、DOC/DOCX
测试方式：构造各类文件内存数据，POST 到 /analyze 接口，验证返回结果
"""
import io
import sys
import time
import base64
import struct
import zlib
import requests
import pandas as pd

BASE_URL = "http://127.0.0.1:5000"
ANALYZE_URL = f"{BASE_URL}/analyze"
PASS = "\033[92m[PASS]\033[0m"
FAIL = "\033[91m[FAIL]\033[0m"
INFO = "\033[94m[INFO]\033[0m"
WARN = "\033[93m[WARN]\033[0m"

results = []


def check_server():
    try:
        r = requests.get(BASE_URL, timeout=5)
        print(f"{INFO} Flask 服务在线，状态码: {r.status_code}")
        return True
    except Exception as e:
        print(f"{FAIL} Flask 服务不可达: {e}")
        print("       请先启动 app.py: python app.py")
        return False


def post_file(name, file_bytes, filename, content_type):
    """发送文件到 /analyze，返回 (ok, data, elapsed)"""
    files = {"image": (filename, io.BytesIO(file_bytes), content_type)}
    t0 = time.time()
    try:
        resp = requests.post(ANALYZE_URL, files=files, timeout=180)
        elapsed = time.time() - t0
        try:
            data = resp.json()
        except Exception:
            return False, {"error": f"响应非 JSON（HTTP {resp.status_code}）: {resp.text[:200]}"}, elapsed
        return True, data, elapsed
    except requests.exceptions.Timeout:
        return False, {"error": "请求超时（>180s）"}, time.time() - t0
    except Exception as e:
        return False, {"error": str(e)}, time.time() - t0


def report(label, ok, data, elapsed):
    status = PASS if ok and "result" in data else FAIL
    print(f"\n{'='*60}")
    print(f"{status} [{label}]  用时: {elapsed:.2f}s")
    if "error" in data:
        print(f"  错误: {data['error']}")
        results.append((label, False, data["error"]))
    elif "result" in data:
        preview = data["result"][:300].replace("\n", " ")
        print(f"  返回（前300字）: {preview}")
        results.append((label, True, "OK"))
    else:
        print(f"  未知响应: {data}")
        results.append((label, False, str(data)))


# ─────────────────────────────────────────────────────────────────
# 1. 构造最小合法 PNG（1×1 红色像素）
# ─────────────────────────────────────────────────────────────────
def make_png_bytes():
    def chunk(name, data):
        c = struct.pack(">I", len(data)) + name + data
        return c + struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = b"\x00\xFF\x00\x00"  # filter=0, R=255, G=0, B=0
    idat = zlib.compress(raw)
    png = b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")
    return png


# ─────────────────────────────────────────────────────────────────
# 2. 构造最小合法 PDF（纯文本，无需 poppler 解析图片层）
#    注：服务端会用 pdf2image 转图，此处用真实简单 PDF
# ─────────────────────────────────────────────────────────────────
def make_pdf_bytes():
    """生成一个包含简单文本的合法 PDF"""
    pdf = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>
stream
BT /F1 12 Tf 100 700 Td (Test PDF Document) Tj ET
endstream
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000346 00000 n 
trailer<</Size 6/Root 1 0 R>>
startxref
441
%%EOF"""
    return pdf


# ─────────────────────────────────────────────────────────────────
# 3. 构造 XLSX（内存生成，openpyxl）
# ─────────────────────────────────────────────────────────────────
def make_xlsx_bytes():
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        pd.DataFrame({
            "收发货人": ["深圳测试贸易有限公司"],
            "监管方式": ["一般贸易"],
            "成交方式": ["FOB"],
            "合计件数": [10],
            "合计毛重(KG)": [60.0],
            "币制": ["USD"],
        }).to_excel(writer, sheet_name="头信息", index=False)
        pd.DataFrame({
            "项号": [1, 2],
            "品名": ["Circuit Board", "Power Supply"],
            "HS编码": ["8534009000", "8504409900"],
            "数量": [100, 50],
            "总价(USD)": [5000.0, 3500.0],
        }).to_excel(writer, sheet_name="货物明细", index=False)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────
# 4. 构造 XLS（xlwt，若未安装则跳过）
# ─────────────────────────────────────────────────────────────────
def make_xls_bytes():
    try:
        import xlwt
        wb = xlwt.Workbook()
        ws = wb.add_sheet("报关单")
        headers = ["项号", "品名", "HS编码", "数量", "总价(USD)"]
        for col, h in enumerate(headers):
            ws.write(0, col, h)
        ws.write(1, 0, 1); ws.write(1, 1, "Circuit Board")
        ws.write(1, 2, "8534009000"); ws.write(1, 3, 100); ws.write(1, 4, 5000.0)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except ImportError:
        return None


# ─────────────────────────────────────────────────────────────────
# 5. 构造 DOCX（python-docx，若未安装则跳过）
# ─────────────────────────────────────────────────────────────────
def make_docx_bytes():
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("测试报关单", 0)
        doc.add_paragraph("收发货人：深圳测试贸易有限公司")
        doc.add_paragraph("监管方式：一般贸易")
        doc.add_paragraph("成交方式：FOB")
        table = doc.add_table(rows=3, cols=5)
        headers = ["项号", "品名", "HS编码", "数量", "总价(USD)"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        row_data = [["1", "Circuit Board", "8534009000", "100", "5000"],
                    ["2", "Power Supply",  "8504409900", "50",  "3500"]]
        for r, row in enumerate(row_data, start=1):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return None


# ─────────────────────────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"\n{'='*60}")
    print("  全格式传输测试  PDF / 图片 / XLSX / XLS / DOCX")
    print(f"{'='*60}")

    if not check_server():
        sys.exit(1)

    # ── 图片测试 ──────────────────────────────────────────────────
    print(f"\n{INFO} [1/5] 测试图片 (PNG)...")
    png = make_png_bytes()
    ok, data, elapsed = post_file("图片(PNG)", png, "test.png", "image/png")
    report("图片(PNG)", ok, data, elapsed)

    # ── XLSX 测试 ─────────────────────────────────────────────────
    print(f"\n{INFO} [2/5] 测试 XLSX...")
    xlsx = make_xlsx_bytes()
    ok, data, elapsed = post_file("Excel(XLSX)", xlsx, "test.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    report("Excel(XLSX)", ok, data, elapsed)

    # ── XLS 测试 ──────────────────────────────────────────────────
    print(f"\n{INFO} [3/5] 测试 XLS...")
    xls_bytes = make_xls_bytes()
    if xls_bytes:
        ok, data, elapsed = post_file("Excel(XLS)", xls_bytes, "test.xls", "application/vnd.ms-excel")
        report("Excel(XLS)", ok, data, elapsed)
    else:
        print(f"{WARN} xlwt 未安装，跳过 XLS 测试（pip install xlwt）")
        results.append(("Excel(XLS)", None, "跳过（xlwt 未安装）"))

    # ── PDF 测试 ──────────────────────────────────────────────────
    print(f"\n{INFO} [4/5] 测试 PDF...")
    pdf = make_pdf_bytes()
    ok, data, elapsed = post_file("PDF", pdf, "test.pdf", "application/pdf")
    report("PDF", ok, data, elapsed)

    # ── DOCX 测试 ─────────────────────────────────────────────────
    print(f"\n{INFO} [5/5] 测试 DOCX...")
    docx_bytes = make_docx_bytes()
    if docx_bytes:
        ok, data, elapsed = post_file("Word(DOCX)", docx_bytes, "test.docx",
                                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        report("Word(DOCX)", ok, data, elapsed)
    else:
        print(f"{WARN} python-docx 未安装，跳过 DOCX 测试（pip install python-docx）")
        results.append(("Word(DOCX)", None, "跳过（python-docx 未安装）"))

    # ── 汇总 ──────────────────────────────────────────────────────
    print(f"\n{'='*60}")
    print("  测试汇总")
    print(f"{'='*60}")
    passed = sum(1 for _, s, _ in results if s is True)
    failed = sum(1 for _, s, _ in results if s is False)
    skipped = sum(1 for _, s, _ in results if s is None)
    for label, status, msg in results:
        if status is True:
            print(f"  {PASS} {label}")
        elif status is False:
            print(f"  {FAIL} {label}  → {msg}")
        else:
            print(f"  {WARN} {label}  → {msg}")
    print(f"\n  通过: {passed}  失败: {failed}  跳过: {skipped}")
    print(f"{'='*60}\n")
    sys.exit(0 if failed == 0 else 1)
